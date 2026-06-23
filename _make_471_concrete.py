# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


RED = RGBColor(255, 0, 0)


def body_text(element):
    return "".join(
        t.text or "" for t in element.iter() if t.tag.split("}")[-1] == "t"
    ).strip()


def add_red_paragraph(doc, text="", style=None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        run.font.color.rgb = RED
    return paragraph._p


def add_red_code_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.color.rgb = RED
    run.font.name = "Consolas"
    return paragraph._p


def add_red_table(doc, headers, rows, table_style):
    table = doc.add_table(rows=1, cols=len(headers))
    if table_style is not None:
        table.style = table_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for column_index, text in enumerate(headers):
        cell = table.rows[0].cells[column_index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(str(text))
        run.bold = True
        run.font.color.rgb = RED
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in rows:
        cells = table.add_row().cells
        for column_index, text in enumerate(row):
            cell = cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(str(text))
            run.font.color.rgb = RED
            if len(str(text)) <= 10:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table._tbl


def insert_after(anchor, elements):
    for element in reversed(elements):
        anchor.addnext(element)


def remove_elements(elements):
    for element in elements:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def main():
    sources = [
        path
        for path in Path(".").glob("*LLM優先順序調整版.docx")
        if not path.name.startswith("~$")
    ]
    if not sources:
        raise RuntimeError("Cannot find LLM優先順序調整版 docx.")
    source_path = max(sources, key=lambda path: path.stat().st_mtime)
    output_path = source_path.with_name(
        source_path.stem.replace("順序調整版", "順序調整具體版") + source_path.suffix
    )

    doc = Document(str(source_path))
    table_style = doc.tables[0].style if doc.tables else None

    body = doc._body._body
    children = list(body)
    start_index = None
    table_caption_index = None
    for index, child in enumerate(children):
        if child.tag.split("}")[-1] != "p":
            continue
        text = body_text(child)
        if text.startswith("4.7.1 問題類型判斷與查詢路由"):
            start_index = index
        if start_index is not None and text.startswith("表 4.9 問題類型"):
            table_caption_index = index
            break
    if start_index is None or table_caption_index is None:
        raise RuntimeError("Cannot locate 4.7.1 replacement range.")

    heading = children[start_index]
    old_range = children[start_index + 1 : table_caption_index]

    body_style = doc.paragraphs[0].style
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("本研究目前採用"):
            body_style = paragraph.style
            break

    new_elements = []
    new_elements.append(
        add_red_paragraph(
            doc,
            "本研究目前採用「LLM 優先、規則回退」的查詢路由設計，其目的不是讓語言模型直接回答問題，而是將自然語言問題轉換為系統可執行的路由資料。實作流程分為五個步驟，如表 4.9 所示。",
            body_style,
        )
    )
    new_elements.append(add_red_paragraph(doc, "表 4.9 GraphRAG 查詢路由處理步驟", body_style))
    new_elements.append(
        add_red_table(
            doc,
            ["步驟", "處理內容", "輸出或控制"],
            [
                ["1", "接收使用者自然語言問題。", "原始 question 字串"],
                ["2", "呼叫本地 LLM 進行問題理解。", "要求只回傳 JSON"],
                ["3", "解析 query_type 與 entities。", "取得檢索類型與查詢線索"],
                ["4", "檢查 JSON 格式與支援類型。", "不合法時改用規則式路由"],
                ["5", "建立 QueryRoute。", "決定 retrieval_modes、use_semantic 與 answer_style"],
            ],
            table_style,
        )
    )
    new_elements.append(
        add_red_paragraph(
            doc,
            "在 LLM 輔助路由階段，系統提示模型只能輸出指定 JSON 結構，且不得產生 Cypher 查詢或自然語言答案。若 LLM 無法連線、回傳內容不是有效 JSON、query_type 不屬於系統支援類型，或必要 entities 無法抽取，系統會回退至規則式判斷。規則式判斷主要依據會議名稱、人名、日期、產品、法規、完成狀態與追蹤語彙等明確線索決定查詢路由。",
            body_style,
        )
    )
    new_elements.append(
        add_red_paragraph(
            doc,
            "LLM 回傳 JSON 範例如下，此範例表示問題被判定為關係查詢，並抽取出人員與法規條件：",
            body_style,
        )
    )
    for line in [
        "{",
        '  "query_type": "relation_lookup",',
        '  "entities": {',
        '    "meeting_hint": "",',
        '    "person_name": "Person_A",',
        '    "date_value": "",',
        '    "product_name": "",',
        '    "regulation_name": "FDA",',
        '    "status": "",',
        '    "keyword": ""',
        "  },",
        '  "confidence": 0.91',
        "}",
    ]:
        new_elements.append(add_red_code_paragraph(doc, line, body_style))
    new_elements.append(
        add_red_paragraph(
            doc,
            "系統解析上述 JSON 後，會將 query_type 對應至預先定義的檢索模式。例如 relation_lookup 對應 relation 模式，後續會查詢 Neo4j 中 Person、Date、Product 或 Regulation 與 MeetingItem 的關係；若 query_type 為 meeting_summary 或 keyword_exploration，系統才會視需要加入 Qdrant 語意檢索。透過此設計，LLM 只負責問題理解與路由輔助，實際資料查詢仍由系統固定流程執行。",
            body_style,
        )
    )

    insert_after(heading, new_elements)
    remove_elements(old_range)

    # Existing query-type table becomes 4.10 because the route-step table is now 4.9.
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("表 4.9 問題類型"):
            for run in paragraph.runs:
                run.text = ""
            run = paragraph.add_run("表 4.10 問題類型、判斷線索與檢索模式對應")
            run.font.color.rgb = RED
        elif paragraph.text.strip().startswith("表 4.10 GraphRAG 問答範圍參數"):
            for run in paragraph.runs:
                run.text = ""
            run = paragraph.add_run("表 4.11 GraphRAG 問答範圍參數")
            run.font.color.rgb = RED
        elif paragraph.text.strip().startswith("表 4.11 GraphRAG 主要 Neo4j 查詢模式"):
            for run in paragraph.runs:
                run.text = ""
            run = paragraph.add_run("表 4.12 GraphRAG 主要 Neo4j 查詢模式")
            run.font.color.rgb = RED

    doc.save(str(output_path))
    print(output_path.resolve())


if __name__ == "__main__":
    main()
