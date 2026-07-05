from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import RGBColor


RED = RGBColor(255, 0, 0)


def set_red_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    run.font.color.rgb = RED


def insert_before(paragraph, text: str):
    new_p = paragraph.insert_paragraph_before("")
    set_red_text(new_p, text)
    return new_p


def main() -> None:
    source = max(
        [p for p in Path(".").glob("*第六章與附錄.docx") if not p.name.startswith("~$")],
        key=lambda p: p.stat().st_mtime,
    )
    doc = Document(source)

    ref_toc = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("參考文獻\t"):
            ref_toc = paragraph
            break
    if ref_toc is None:
        raise RuntimeError("TOC reference line not found.")

    # Avoid duplicate insertions if this script is rerun.
    toc_text = "\n".join(p.text for p in doc.paragraphs[:180])
    if "第六章 結論與未來發展\t51" not in toc_text:
        insert_before(ref_toc, "第六章 結論與未來發展\t51")
        insert_before(ref_toc, "6.1 研究結論\t51")
        insert_before(ref_toc, "6.2 研究限制\t52")
        insert_before(ref_toc, "6.3 未來發展\t52")
    set_red_text(ref_toc, "參考文獻\t54")
    if "附錄A GraphRAG 測試問題與判定結果\t57" not in toc_text:
        insert_before(ref_toc, "附錄A GraphRAG 測試問題與判定結果\t57")

    out = source.with_name(source.stem + "_目錄更新.docx")
    doc.save(out)
    print(out.resolve())


if __name__ == "__main__":
    main()
