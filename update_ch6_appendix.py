from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


RED = RGBColor(255, 0, 0)


def latest_source() -> Path:
    files = [p for p in Path(".").glob("*.docx") if not p.name.startswith("~$")]
    candidates = [p for p in files if "5.6" in p.name or p.name == "paper_table56_updated.docx"]
    if not candidates:
        raise RuntimeError("No Table 5.6 updated DOCX found.")
    return max(candidates, key=lambda p: p.stat().st_mtime)


def format_para(paragraph, kind: str = "body") -> None:
    if kind == "chapter":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.color.rgb = RED
            run.bold = True
            run.font.size = Pt(16)
    elif kind == "section":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(4)
        for run in paragraph.runs:
            run.font.color.rgb = RED
            run.bold = True
            run.font.size = Pt(14)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(6)
        for run in paragraph.runs:
            run.font.color.rgb = RED
            run.font.size = Pt(12)


def add_red_paragraph(doc: Document, text: str, kind: str = "body"):
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)
    format_para(paragraph, kind)
    return paragraph


def set_cell_text(cell, text: str, *, align=None, bold: bool = False, size: float | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    run.font.color.rgb = RED
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def remove_table(table) -> None:
    table._tbl.getparent().remove(table._tbl)


def move_before(element, reference_paragraph) -> None:
    reference_paragraph._p.addprevious(element)


def find_table_56_1(doc: Document):
    for table in doc.tables:
        if len(table.columns) != 3 or not table.rows:
            continue
        header = [cell.text.strip().replace("\n", "") for cell in table.rows[0].cells]
        if header == ["問題類型", "測試問題", "判定"]:
            return table
    raise RuntimeError("Table 5.6-1 not found.")


def find_table_56(doc: Document):
    for table in doc.tables:
        if len(table.columns) != 7 or not table.rows:
            continue
        header = [cell.text.strip().replace("\n", "") for cell in table.rows[0].cells]
        if header[0] == "問題類型" and header[1] == "檢索來源":
            return table
    raise RuntimeError("Table 5.6 not found.")


def chapter6_content() -> list[tuple[str, str]]:
    return [
        ("chapter", "第六章 結論與未來發展"),
        ("section", "6.1 研究結論"),
        (
            "body",
            "本研究以企業會議記錄為例，提出一套結合結構化文件解析、知識圖譜建構、向量檢索與來源限制回答生成之 GraphRAG 檢索系統。此系統將原本以 PDF 形式保存的會議記錄轉換為 Document、Meeting 與 MeetingItem 三個層級，並以 MeetingItem 作為最小證據單位，使會議內容、負責人、預計日期、實際完成日與追蹤結果等欄位能被系統化保存與查詢。相較於僅以整份文件或文字片段為單位的檢索方式，本研究強調回答必須能回到具體會議項目與來源欄位，以提升企業會議記錄查詢的可驗證性。",
        ),
        (
            "body",
            "在系統實作方面，本研究完成 PDF 文字層解析、MongoDB 結構化儲存、Neo4j 知識圖譜建構、Qdrant 向量索引與 GraphRAG 問答流程整合。實驗資料共匯入 10 份去識別化企業會議記錄，其中 9 份具文字層內容並成功解析為 9 場會議與 116 筆 MeetingItem；另 1 份文件因需 OCR 處理，標記為 needs_ocr。知識圖譜部分建立 13 類節點共 893 個與 23 類關係共 4,946 筆，向量索引則以 116 筆 MeetingItem 建立 Qdrant point，使每一筆已解析會議項目皆具備語意檢索入口。",
        ),
        (
            "body",
            "在問答測試方面，本研究重新依目前去識別化資料設計 18 題 GraphRAG 測試問題，涵蓋責任歸屬、時程狀態、關鍵詞探索、產品或法規關聯、開放模糊／語意相似與資料不足等類型。測試結果顯示，責任歸屬型、時程狀態型、開放模糊型／語意相似型與資料不足型皆可回到來源證據或正確拒答；關鍵詞探索型雖能召回相關 MeetingItem，但因其關係較適合作探索，不宜直接解讀為正式責任、決議或因果關係，因此採較保守之部分正確判定。整體而言，系統能在多數查詢情境中維持回答、來源列表與圖譜證據之對應關係。",
        ),
        (
            "body",
            "本研究之主要貢獻可歸納為四點。第一，建立以會議項目為核心的結構化文件模型，使會議記錄能由文件層級下鑽至具體事項。第二，將負責人、日期、產品、法規與關鍵詞轉換為知識圖譜關係，使原本分散於會議記錄中的脈絡可以被查詢與展開。第三，整合向量檢索與圖譜查詢，使系統能同時支援語意相似查詢與欄位事實回溯。第四，透過 evidence_id、meeting_id、item_id 與 item_no 等來源資訊限制回答生成，降低語言模型在企業責任歸屬、時程判斷與決議查詢中的無來源生成風險。",
        ),
        ("section", "6.2 研究限制"),
        (
            "body",
            "本研究仍受到資料格式與資料規模限制。實驗資料以固定格式且具文字層之企業會議記錄 PDF 為主，對於掃描型 PDF、版面差異較大的文件或表格結構不穩定的會議記錄，仍需加入 OCR、版面偵測與解析前文件品質檢查，才能擴大系統適用範圍。因此，本研究結果較適合說明固定格式企業會議記錄導入結構化解析與 GraphRAG 檢索流程之可行性，尚不能直接推論至所有非結構化企業文件。",
        ),
        (
            "body",
            "其次，知識圖譜關係品質仍依賴欄位解析與實體抽取結果。RESPONSIBLE_BY、HAS_PLANNED_DATE 與 HAS_COMPLETED_DATE 等欄位型關係可直接支援責任與時程查詢，但 MENTIONS 與 CO_OCCURS_WITH 等內容型關係較適合作為探索性輔助，不宜直接解讀為正式責任、因果關係或決議依據。若未在回答階段區分強關係與弱關係，系統可能將關鍵詞共現或語意相近誤作事實判斷。",
        ),
        (
            "body",
            "此外，向量檢索在本研究中主要扮演候選證據召回角色，不能單獨作為事實正確性的依據。摘要型或開放模糊型問題雖可透過向量檢索召回較多候選 MeetingItem，但候選範圍可能偏大，仍需透過 evidence selector、meeting_hint 或欄位條件進一步縮小證據集合。未來若要將本系統用於正式決策支援，仍需建立更完整的人工標註基準、Top-k 命中率、來源回溯正確率與回答事實一致性評估。",
        ),
        ("section", "6.3 未來發展"),
        (
            "body",
            "未來研究可優先強化文件解析與資料前處理能力。針對 needs_ocr 文件，可加入 OCR、表格版面偵測與解析前品質檢查，使掃描型 PDF 與版面差異較大的文件能被分流處理。若能在解析階段更穩定地辨識會議名稱、項次、負責人、日期與追蹤結果，後續 MongoDB 結構化欄位、Neo4j 圖譜關係與 Qdrant 向量索引的品質也將同步提升。",
        ),
        (
            "body",
            "第二，未來可強化問題路由與證據選擇機制。本研究目前採 LLM 優先並搭配規則式 fallback 的路由方式，能在 LLM 逾時或失敗時維持基本查詢能力。後續可進一步建立可評估的 query_type 標註資料，檢查 LLM 路由、規則式路由與混合路由在不同問題類型上的差異，並針對摘要型、關鍵詞探索型與複合條件查詢設計更嚴格的 evidence selector，以降低候選證據過寬或弱關係過度解讀的風險。",
        ),
        (
            "body",
            "第三，未來可建立更完整的量化評估與回歸測試流程。當資料經去識別化、重新解析或重新索引後，item_id 與 meeting_id 可能重新產生，因此 golden cases 應與新資料同步更新。後續可針對各問題類型建立 expected evidence、Top-k 命中率、來源回溯正確率、欄位依據正確率與拒答正確率等指標，並將其整合為固定回歸測試，以利系統版本更新後持續檢查回答品質。",
        ),
        (
            "body",
            "最後，本研究系統可延伸至更廣泛的企業知識管理情境。除會議記錄外，企業內部常見的設計審查紀錄、專案追蹤表、品質異常報告與法規送件文件，也具有欄位事實、跨文件關聯與來源查核需求。未來若能建立跨文件類型的資料模型與圖譜 schema，並結合權限控管、使用者回饋與互動式圖譜視覺化，將有助於將 GraphRAG 從單一文件檢索工具擴展為企業內部可追溯、可解釋的知識查詢平台。",
        ),
    ]


def main() -> None:
    src = latest_source()
    out = src.with_name(src.stem + "_第六章與附錄.docx")
    doc = Document(src)

    source_labels = {
        "責任歸屬型": "MongoDB 結構化欄位、Neo4j 知識圖譜",
        "時程狀態型": "MongoDB 結構化欄位、Neo4j 知識圖譜",
        "關鍵詞探索型": "Neo4j 知識圖譜、Qdrant 向量、MongoDB 回查",
        "產品或法規關聯型": "Neo4j 知識圖譜、MongoDB 結構化欄位回查",
        "開放模糊型／語意相似型": "Qdrant 向量、MongoDB 結構化欄位回查",
        "資料不足型": "GraphRAG 證據集合檢查",
    }
    table_56 = find_table_56(doc)
    for row in table_56.rows[1:]:
        category = row.cells[0].text.strip().replace("\n", "")
        if category in source_labels:
            set_cell_text(row.cells[1], source_labels[category], align=WD_ALIGN_PARAGRAPH.LEFT)

    table_56_1 = find_table_56_1(doc)
    appendix_table_xml = deepcopy(table_56_1._tbl)
    remove_table(table_56_1)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("表 5.6-1"):
            remove_paragraph(paragraph)
            break

    caption_57 = next(p for p in doc.paragraphs if p.text.strip().startswith("表 5.7"))
    main_note = add_red_paragraph(
        doc,
        "完整測試題目與逐題判定結果移列附錄A，本節以問題類型統計與代表案例呈現主要結果。",
        "body",
    )
    move_before(main_note._p, caption_57)

    reference_paragraph = next(p for p in doc.paragraphs if p.text.strip() == "參考文獻")
    chapter_elements = []
    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)
    chapter_elements.append(page_break._p)
    for kind, text in chapter6_content():
        chapter_elements.append(add_red_paragraph(doc, text, kind)._p)
    page_break_refs = doc.add_paragraph()
    page_break_refs.add_run().add_break(WD_BREAK.PAGE)
    chapter_elements.append(page_break_refs._p)
    for element in chapter_elements:
        move_before(element, reference_paragraph)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_red_paragraph(doc, "附錄A GraphRAG 測試問題與判定結果", "chapter")
    add_red_paragraph(
        doc,
        "本附錄列出第五章表 5.6 統計所使用之 18 題測試問題與人工判定結果。測試題依目前去識別化資料重新設計，作為 GraphRAG 問答評估之逐題紀錄。",
        "body",
    )
    caption = add_red_paragraph(doc, "附表 A-1 GraphRAG 測試問題與判定結果", "body")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc._body._element.append(appendix_table_xml)
    appendix_table = doc.tables[-1]
    set_table_widths(appendix_table, [1800, 6200, 1200])
    for row in appendix_table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.color.rgb = RED
                    run.font.size = Pt(9.5)

    doc.save(out)
    print(f"SRC={src.name}")
    print(f"OUT={out.resolve()}")


if __name__ == "__main__":
    main()
